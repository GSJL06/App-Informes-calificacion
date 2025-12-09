#!/usr/bin/env python
"""
Script para analizar y refactorizar plantillas DOCX.
Detecta contenido hardcodeado (listas, tablas) y lo convierte a placeholders dinámicos.

Uso:
    python refactorizar_plantillas.py --analizar
    python refactorizar_plantillas.py --auto-refactor
    python refactorizar_plantillas.py --plantilla templates/plantilla_desempeno.docx --analizar
"""
import sys
sys.path.insert(0, 'src')

import argparse
import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class TemplateAnalyzer:
    """Analyzes DOCX templates for hardcoded content."""
    
    PLACEHOLDER_PATTERN = r'\{\{([a-zA-Z0-9_]+)\}\}'
    
    def __init__(self, template_path: Path):
        self.template_path = template_path
        self.doc = Document(template_path)
        self.pattern = re.compile(self.PLACEHOLDER_PATTERN)
        
    def analyze(self) -> Dict[str, Any]:
        """Full analysis of template."""
        return {
            'template': str(self.template_path),
            'existing_placeholders': self._find_placeholders(),
            'bullet_lists': self._find_bullet_lists(),
            'tables_with_hardcoded': self._find_hardcoded_tables(),
            'statistics': self._get_statistics()
        }
    
    def _find_placeholders(self) -> List[str]:
        """Find all existing placeholders."""
        placeholders = set()
        
        for para in self.doc.paragraphs:
            placeholders.update(self.pattern.findall(para.text))
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        placeholders.update(self.pattern.findall(para.text))
        
        for section in self.doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    placeholders.update(self.pattern.findall(para.text))
            if section.footer:
                for para in section.footer.paragraphs:
                    placeholders.update(self.pattern.findall(para.text))
        
        return sorted(placeholders)
    
    def _find_bullet_lists(self) -> List[Dict]:
        """Find bullet lists with potential hardcoded content."""
        lists = []
        current_list = []
        list_start_idx = None
        
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else ''
            
            is_bullet = (
                text.startswith('- ') or 
                text.startswith('• ') or
                'List' in style_name or
                'Bullet' in style_name or
                (para._element.pPr is not None and 
                 para._element.pPr.numPr is not None)
            )
            
            if is_bullet and text:
                if not current_list:
                    list_start_idx = idx
                current_list.append({
                    'index': idx,
                    'text': text,
                    'has_placeholder': bool(self.pattern.search(text))
                })
            else:
                if len(current_list) >= 2:
                    has_hardcoded = any(not item['has_placeholder'] for item in current_list)
                    lists.append({
                        'start_index': list_start_idx,
                        'items': current_list,
                        'count': len(current_list),
                        'has_hardcoded_items': has_hardcoded,
                        'needs_refactor': has_hardcoded
                    })
                current_list = []
                list_start_idx = None
        
        if len(current_list) >= 2:
            has_hardcoded = any(not item['has_placeholder'] for item in current_list)
            lists.append({
                'start_index': list_start_idx,
                'items': current_list,
                'count': len(current_list),
                'has_hardcoded_items': has_hardcoded,
                'needs_refactor': has_hardcoded
            })
        
        return lists
    
    def _find_hardcoded_tables(self) -> List[Dict]:
        """Find tables with hardcoded data."""
        tables_info = []
        
        for table_idx, table in enumerate(self.doc.tables):
            table_data = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'hardcoded_cells': [],
                'placeholder_cells': [],
                'needs_refactor': False
            }
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    
                    has_placeholder = bool(self.pattern.search(cell_text))
                    
                    cell_info = {
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell_text[:100],
                        'has_placeholder': has_placeholder
                    }
                    
                    if has_placeholder:
                        table_data['placeholder_cells'].append(cell_info)
                    elif row_idx > 0:  # Skip header row
                        # Check if looks like data (not a label)
                        if self._looks_like_data(cell_text):
                            table_data['hardcoded_cells'].append(cell_info)
                            table_data['needs_refactor'] = True
            
            if table_data['hardcoded_cells'] or table_data['placeholder_cells']:
                tables_info.append(table_data)
        
        return tables_info
    
    def _looks_like_data(self, text: str) -> bool:
        """Heuristic to detect if text looks like hardcoded data."""
        # Phone numbers
        if re.match(r'^[\d\s\-\+]{7,}$', text):
            return True
        # Email
        if '@' in text and '.' in text:
            return True
        # Specific names (capitalized words)
        if re.match(r'^[A-ZÁÉÍÓÚ][a-záéíóú]+ [A-ZÁÉÍÓÚ][a-záéíóú]+', text):
            return True
        # Codes like PRF196
        if re.match(r'^[A-Z]{2,4}\d{2,}', text):
            return True
        return False
    
    def _get_statistics(self) -> Dict:
        """Get document statistics."""
        return {
            'sections': len(self.doc.sections),
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables)
        }


class TemplateRefactorer:
    """Refactors templates to use dynamic placeholders."""
    
    def __init__(self, template_path: Path):
        self.template_path = template_path
        self.doc = Document(template_path)
        self.changes = []
        
    def create_backup(self) -> Path:
        """Create backup of template."""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_path = self.template_path.with_suffix(f'.backup.{timestamp}.docx')
        shutil.copy2(self.template_path, backup_path)
        logger.info(f"Backup created: {backup_path}")
        return backup_path
    
    def refactor_bullet_list(self, list_info: Dict, placeholder_name: str) -> bool:
        """Replace bullet list with single placeholder."""
        if not list_info['items']:
            return False
        
        start_idx = list_info['start_index']
        first_para = self.doc.paragraphs[start_idx]
        
        # Store original formatting
        original_style = first_para.style
        
        # Set placeholder in first item
        first_para.clear()
        run = first_para.add_run(f"{{{{{placeholder_name}}}}}")
        
        # Remove other items (mark for deletion)
        for item in list_info['items'][1:]:
            para = self.doc.paragraphs[item['index']]
            para.clear()
        
        self.changes.append({
            'type': 'bullet_list',
            'placeholder': placeholder_name,
            'original_items': [i['text'] for i in list_info['items']]
        })
        
        return True
    
    def refactor_table_column(self, table_idx: int, col_idx: int, 
                              placeholder_base: str) -> bool:
        """Add placeholder pattern to table column."""
        table = self.doc.tables[table_idx]
        
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:  # Skip header
                continue
            
            cell = row.cells[col_idx]
            original_text = cell.text.strip()
            
            if original_text and not re.search(r'\{\{.*\}\}', original_text):
                # Replace with placeholder
                placeholder = f"{{{{{placeholder_base}_{row_idx}}}}}"
                for para in cell.paragraphs:
                    para.clear()
                    para.add_run(placeholder)
                
                self.changes.append({
                    'type': 'table_cell',
                    'table': table_idx,
                    'row': row_idx,
                    'col': col_idx,
                    'placeholder': placeholder,
                    'original': original_text
                })
        
        return True
    
    def save(self, output_path: Optional[Path] = None) -> Path:
        """Save refactored template."""
        save_path = output_path or self.template_path
        self.doc.save(save_path)
        logger.info(f"Saved refactored template: {save_path}")
        return save_path
    
    def generate_example_json(self) -> Dict:
        """Generate example JSON with new placeholders."""
        example = {}
        
        for change in self.changes:
            if change['type'] == 'bullet_list':
                example[change['placeholder']] = change['original_items']
            elif change['type'] == 'table_cell':
                example[change['placeholder'].strip('{}')] = change['original']
        
        return example


def analyze_all_templates(templates_dir: Path) -> Dict:
    """Analyze all templates in directory."""
    results = {}
    
    for template_file in templates_dir.glob('plantilla_*.docx'):
        if '.backup.' in template_file.name:
            continue
        
        logger.info(f"Analyzing: {template_file.name}")
        analyzer = TemplateAnalyzer(template_file)
        results[template_file.name] = analyzer.analyze()
    
    return results


def print_analysis_report(analysis: Dict):
    """Print formatted analysis report."""
    print("\n" + "=" * 70)
    print("TEMPLATE ANALYSIS REPORT")
    print("=" * 70)
    
    for template_name, data in analysis.items():
        print(f"\n📄 {template_name}")
        print("-" * 50)
        
        # Existing placeholders
        placeholders = data['existing_placeholders']
        print(f"\n  ✓ Existing placeholders ({len(placeholders)}):")
        for ph in placeholders[:10]:
            print(f"    - {{{{{ph}}}}}")
        if len(placeholders) > 10:
            print(f"    ... and {len(placeholders) - 10} more")
        
        # Bullet lists
        lists = data['bullet_lists']
        needs_refactor = [l for l in lists if l['needs_refactor']]
        print(f"\n  📋 Bullet lists found: {len(lists)}")
        if needs_refactor:
            print(f"    ⚠️  Lists needing refactor: {len(needs_refactor)}")
            for lst in needs_refactor:
                print(f"      - {lst['count']} items starting at paragraph {lst['start_index']}")
                for item in lst['items'][:3]:
                    print(f"        • {item['text'][:60]}...")
        
        # Tables
        tables = data['tables_with_hardcoded']
        tables_needing_refactor = [t for t in tables if t['needs_refactor']]
        print(f"\n  📊 Tables with data: {len(tables)}")
        if tables_needing_refactor:
            print(f"    ⚠️  Tables needing refactor: {len(tables_needing_refactor)}")
            for tbl in tables_needing_refactor:
                print(f"      - Table {tbl['index']}: {tbl['rows']}x{tbl['cols']}")
                print(f"        Hardcoded cells: {len(tbl['hardcoded_cells'])}")
                for cell in tbl['hardcoded_cells'][:3]:
                    print(f"          [{cell['row']},{cell['col']}]: {cell['text'][:40]}...")
        
        # Statistics
        stats = data['statistics']
        print(f"\n  📈 Statistics:")
        print(f"    - Sections: {stats['sections']}")
        print(f"    - Paragraphs: {stats['paragraphs']}")
        print(f"    - Tables: {stats['tables']}")
    
    print("\n" + "=" * 70)


def auto_refactor_template(template_path: Path, dry_run: bool = False) -> Dict:
    """Auto-refactor a single template."""
    analyzer = TemplateAnalyzer(template_path)
    analysis = analyzer.analyze()
    
    if dry_run:
        return {'analysis': analysis, 'changes': [], 'dry_run': True}
    
    refactorer = TemplateRefactorer(template_path)
    backup_path = refactorer.create_backup()
    
    # Refactor bullet lists
    for idx, lst in enumerate(analysis['bullet_lists']):
        if lst['needs_refactor']:
            placeholder_name = f"lista_dinamica_{idx}"
            refactorer.refactor_bullet_list(lst, placeholder_name)
    
    # Save changes
    refactorer.save()
    
    # Generate example JSON
    example_json = refactorer.generate_example_json()
    
    return {
        'analysis': analysis,
        'changes': refactorer.changes,
        'backup': str(backup_path),
        'example_json': example_json
    }


def main():
    parser = argparse.ArgumentParser(
        description='Analyze and refactor DOCX templates for dynamic content'
    )
    
    parser.add_argument(
        '--analizar',
        action='store_true',
        help='Analyze templates without modifying'
    )
    
    parser.add_argument(
        '--auto-refactor',
        action='store_true',
        help='Auto-refactor templates with backup'
    )
    
    parser.add_argument(
        '--plantilla',
        type=str,
        help='Specific template to process'
    )
    
    parser.add_argument(
        '--output-json',
        type=str,
        help='Output JSON file for analysis results'
    )
    
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Show what would be changed without modifying'
    )
    
    args = parser.parse_args()
    
    templates_dir = Path('templates')
    
    if args.analizar:
        if args.plantilla:
            template_path = Path(args.plantilla)
            if not template_path.exists():
                print(f"❌ Template not found: {template_path}")
                sys.exit(1)
            analyzer = TemplateAnalyzer(template_path)
            analysis = {template_path.name: analyzer.analyze()}
        else:
            analysis = analyze_all_templates(templates_dir)
        
        print_analysis_report(analysis)
        
        if args.output_json:
            with open(args.output_json, 'w', encoding='utf-8') as f:
                json.dump(analysis, f, indent=2, ensure_ascii=False)
            print(f"\n✅ Analysis saved to: {args.output_json}")
    
    elif args.auto_refactor:
        if args.plantilla:
            template_path = Path(args.plantilla)
            if not template_path.exists():
                print(f"❌ Template not found: {template_path}")
                sys.exit(1)
            
            result = auto_refactor_template(template_path, dry_run=args.dry_run)
            
            if args.dry_run:
                print("\n🔍 DRY RUN - No changes made")
                print_analysis_report({template_path.name: result['analysis']})
            else:
                print(f"\n✅ Template refactored: {template_path}")
                print(f"   Backup: {result['backup']}")
                print(f"   Changes: {len(result['changes'])}")
                
                if result['example_json']:
                    example_path = template_path.with_suffix('.example.json')
                    with open(example_path, 'w', encoding='utf-8') as f:
                        json.dump(result['example_json'], f, indent=2, ensure_ascii=False)
                    print(f"   Example JSON: {example_path}")
        else:
            print("❌ Please specify --plantilla for auto-refactor")
            sys.exit(1)
    
    else:
        parser.print_help()


if __name__ == '__main__':
    main()