"""
Tests para el módulo ImageReplacer.
"""
import sys
import os
import tempfile
import shutil
from pathlib import Path
from unittest import TestCase, main
from unittest.mock import Mock, patch, MagicMock

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))

from docx import Document
from core.image_replacer import ImageReplacer, ImageInfo, replace_images_in_document


class TestImageInfo(TestCase):
    """Tests para la clase ImageInfo."""
    
    def test_image_info_creation(self):
        """Test creación de ImageInfo."""
        info = ImageInfo(
            rel_id='rId5',
            target='media/image1.png',
            location='header',
            section_idx=0,
            index=0,
            width_emu=914400,
            height_emu=457200
        )
        
        self.assertEqual(info.rel_id, 'rId5')
        self.assertEqual(info.location, 'header')
        self.assertEqual(info.width_emu, 914400)
        self.assertEqual(info.height_emu, 457200)
    
    def test_width_inches_conversion(self):
        """Test conversión de EMUs a pulgadas."""
        info = ImageInfo(
            rel_id='rId1',
            target='media/image1.png',
            location='body',
            width_emu=914400,  # 1 inch
            height_emu=1828800  # 2 inches
        )
        
        self.assertAlmostEqual(info.width_inches, 1.0, places=5)
        self.assertAlmostEqual(info.height_inches, 2.0, places=5)
    
    def test_width_inches_none(self):
        """Test conversión cuando no hay dimensiones."""
        info = ImageInfo(
            rel_id='rId1',
            target='media/image1.png',
            location='body'
        )
        
        self.assertIsNone(info.width_inches)
        self.assertIsNone(info.height_inches)
    
    def test_to_dict(self):
        """Test conversión a diccionario."""
        info = ImageInfo(
            rel_id='rId5',
            target='media/image1.png',
            location='header',
            section_idx=0,
            index=1,
            width_emu=914400,
            height_emu=457200,
            is_inline=True
        )
        
        result = info.to_dict()
        
        self.assertEqual(result['rel_id'], 'rId5')
        self.assertEqual(result['location'], 'header')
        self.assertEqual(result['section_idx'], 0)
        self.assertEqual(result['index'], 1)
        self.assertTrue(result['is_inline'])


class TestImageReplacer(TestCase):
    """Tests para la clase ImageReplacer."""
    
    def setUp(self):
        """Configuración inicial para cada test."""
        self.test_dir = tempfile.mkdtemp()
        self.test_doc_path = os.path.join(self.test_dir, 'test.docx')
        
        # Crear documento de prueba simple
        doc = Document()
        doc.add_paragraph("Test document")
        doc.save(self.test_doc_path)
    
    def tearDown(self):
        """Limpieza después de cada test."""
        shutil.rmtree(self.test_dir, ignore_errors=True)
    
    def test_replacer_initialization(self):
        """Test inicialización del replacer."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        self.assertIsNotNone(replacer.document)
        self.assertIsInstance(replacer._image_cache, dict)
    
    def test_get_all_images_info_empty(self):
        """Test obtener info de imágenes en documento vacío."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        info = replacer.get_all_images_info()
        
        self.assertIn('headers', info)
        self.assertIn('footers', info)
        self.assertIn('body', info)
        self.assertEqual(len(info['body']), 0)
    
    def test_get_summary_empty(self):
        """Test resumen de documento sin imágenes."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        summary = replacer.get_summary()
        
        self.assertEqual(summary['total'], 0)
        self.assertEqual(summary['total_headers'], 0)
        self.assertEqual(summary['total_body'], 0)
        self.assertEqual(summary['total_footers'], 0)
    
    def test_replace_header_image_invalid_section(self):
        """Test reemplazo con sección inválida."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        result = replacer.replace_header_image(
            section_idx=99,
            new_image_path='fake.png'
        )
        
        self.assertFalse(result)
    
    def test_replace_header_image_file_not_found(self):
        """Test reemplazo con archivo inexistente."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        result = replacer.replace_header_image(
            section_idx=0,
            new_image_path='nonexistent.png'
        )
        
        self.assertFalse(result)
    
    def test_replace_body_image_invalid_index(self):
        """Test reemplazo de imagen con índice inválido."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        result = replacer.replace_body_image_by_index(
            image_index=99,
            new_image_path='fake.png'
        )
        
        self.assertFalse(result)
    
    def test_replace_images_batch_empty(self):
        """Test reemplazo en lote vacío."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        results = replacer.replace_images_batch({})
        
        self.assertEqual(len(results), 0)
    
    def test_replace_images_batch_invalid_location(self):
        """Test reemplazo en lote con ubicación inválida."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        results = replacer.replace_images_batch({
            'invalid_0': 'fake.png'
        })
        
        self.assertFalse(results['invalid_0'])
    
    def test_get_header_images_info_empty(self):
        """Test obtener info de headers sin imágenes."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        info = replacer.get_header_images_info(section_idx=0)
        
        self.assertEqual(len(info), 0)
    
    def test_get_footer_images_info_empty(self):
        """Test obtener info de footers sin imágenes."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        info = replacer.get_footer_images_info(section_idx=0)
        
        self.assertEqual(len(info), 0)
    
    def test_get_body_images_info_empty(self):
        """Test obtener info del body sin imágenes."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        info = replacer.get_body_images_info()
        
        self.assertEqual(len(info), 0)
    
    def test_get_image_dimensions_not_found(self):
        """Test obtener dimensiones de imagen inexistente."""
        doc = Document(self.test_doc_path)
        replacer = ImageReplacer(doc)
        
        dims = replacer.get_image_dimensions('body', 0)
        
        self.assertIsNone(dims)


class TestReplaceImagesInDocument(TestCase):
    """Tests para la función de conveniencia."""
    
    def setUp(self):
        """Configuración inicial."""
        self.test_dir = tempfile.mkdtemp()
        self.test_doc_path = os.path.join(self.test_dir, 'test.docx')
        self.output_path = os.path.join(self.test_dir, 'output.docx')
        
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(self.test_doc_path)
    
    def tearDown(self):
        """Limpieza."""
        shutil.rmtree(self.test_dir, ignore_errors=True)
    
    def test_replace_images_in_document_empty(self):
        """Test función con reemplazos vacíos."""
        result = replace_images_in_document(
            self.test_doc_path,
            self.output_path,
            {}
        )
        
        self.assertTrue(result)
        self.assertTrue(os.path.exists(self.output_path))
    
    def test_replace_images_in_document_invalid_path(self):
        """Test función con ruta inválida."""
        result = replace_images_in_document(
            'nonexistent.docx',
            self.output_path,
            {}
        )
        
        self.assertFalse(result)


class TestImageReplacerWithRealDocument(TestCase):
    """Tests con documentos reales del proyecto."""
    
    @classmethod
    def setUpClass(cls):
        """Verificar si existen documentos de prueba."""
        cls.test_docs_dir = Path('test_documents')
        cls.templates_dir = Path('templates')
        cls.has_test_docs = cls.test_docs_dir.exists()
        cls.has_templates = cls.templates_dir.exists()
    
    def test_analyze_real_document(self):
        """Test análisis de documento real si existe."""
        if not self.has_test_docs:
            self.skipTest("No hay documentos de prueba disponibles")
        
        # Buscar cualquier documento .docx
        docs = list(self.test_docs_dir.glob('*.docx'))
        if not docs:
            self.skipTest("No hay archivos .docx en test_documents")
        
        doc = Document(docs[0])
        replacer = ImageReplacer(doc)
        summary = replacer.get_summary()
        
        self.assertIn('total', summary)
        self.assertIn('details', summary)
    
    def test_analyze_template(self):
        """Test análisis de plantilla si existe."""
        if not self.has_templates:
            self.skipTest("No hay plantillas disponibles")
        
        templates = list(self.templates_dir.glob('*.docx'))
        if not templates:
            self.skipTest("No hay archivos .docx en templates")
        
        doc = Document(templates[0])
        replacer = ImageReplacer(doc)
        summary = replacer.get_summary()
        
        self.assertIsInstance(summary['total'], int)


if __name__ == '__main__':
    main()