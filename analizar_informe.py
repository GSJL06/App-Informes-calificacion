"""
Script para analizar los 4 informes de telemetría
"""
import sys
sys.path.insert(0, 'src')
from docx import Document
from pathlib import Path
import re

# Lista de informes a analizar
informes = [
    'test_documents/Informe_de_calificación_de_desempeño_ejemplo_formato_actual.docx',
    'test_documents/Informe_de_calificación_de_diseño_ejemplo_formato_actual.docx',
    'test_documents/Informe_de_calificación_de_instalación_ejemplo_formato_actual.docx',
    'test_documents/Informe_de_calificación_de_operacion_ejemplo_formato_actual.docx',
]

for informe_path in informes:
    nombre = Path(informe_path).stem.replace('Informe_de_calificación_de_', '').replace('_ejemplo_formato_actual', '')

    print('\n' + '='*70)
    print(f'📄 INFORME: {nombre.upper()}')
    print('='*70)

    doc = Document(informe_path)

    # Estadísticas básicas
    print(f'\n📊 ESTADÍSTICAS:')
    print(f'   Secciones: {len(doc.sections)}')
    print(f'   Tablas: {len(doc.tables)}')
    print(f'   Párrafos: {len(doc.paragraphs)}')

    # Imágenes en el documento
    img_count = sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)
    print(f'   Imágenes en cuerpo: {img_count}')

    # Imágenes en headers
    header_imgs = []
    for section in doc.sections:
        if section.header:
            for rel in section.header.part.rels.values():
                if 'image' in rel.reltype:
                    header_imgs.append(rel.target_ref)
    print(f'   Imágenes en header: {len(header_imgs)} - {header_imgs}')

    # Contenido principal (primeros párrafos importantes)
    print(f'\n📝 DATOS PRINCIPALES:')
    for para in doc.paragraphs[:15]:
        text = para.text.strip()
        if text and ':' in text:
            print(f'   {text[:80]}')

    # Tablas resumen
    print(f'\n📋 ESTRUCTURA DE TABLAS:')
    for i, table in enumerate(doc.tables[:4]):
        first_cell = table.rows[0].cells[0].text[:40].replace('\n', ' ') if table.rows else ''
        print(f'   Tabla {i}: {len(table.rows)}x{len(table.columns)} - "{first_cell}..."')

print('\n' + '='*70)
print('ANÁLISIS COMPLETADO')
print('='*70)
