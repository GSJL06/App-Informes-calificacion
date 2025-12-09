"""
Script para convertir los informes de ejemplo en plantillas con placeholders.
Reemplaza datos específicos por variables {{placeholder}} y genera
un archivo de configuración JSON con información de cada plantilla.
"""
import sys
sys.path.insert(0, 'src')
from docx import Document
from pathlib import Path
import re
import json
from core.image_replacer import ImageReplacer

# Mapeo de datos específicos a placeholders
REPLACEMENTS = {
    # Datos del establecimiento
    "Hospital Pablo Tobón Uribe": "{{nombre_establecimiento}}",
    "Calle 78 B # 69 - 240": "{{direccion}}",
    "Calle 78 B # 69 – 240": "{{direccion}}",  # Variante con guión diferente
    "Calle 78B # 69 - 240": "{{direccion}}",
    
    # Títulos de equipos/sistemas
    "Sistema de telemetría": "{{titulo}}",
    "Sistema de monitoreo": "{{titulo}}",
    
    # Fechas
    "09/07/2025": "{{fecha_calificacion}}",
    "02/07/2025": "{{fecha_calificacion}}",
    "03/07/2025": "{{fecha_calificacion}}",
    "18/07/2025": "{{fecha_firma}}",
    "07/07/2025": "{{fecha_calificacion}}",
    
    # Números de informe
    "N° 3": "{{numero_informe}}",
    "N° 1": "{{numero_informe}}",
    "N° 2": "{{numero_informe}}",
    "N° 4": "{{numero_informe}}",
    
    # Contactos y responsables específicos
    "Maria Alejandra Zapata Ch": "{{nombre_realizo}}",
    "Rubén Darío Tabares Muñoz": "{{nombre_verifico}}",
    "yahoyos@hptu.org.co": "{{email_contacto_1}}",
    "laangarita@hptu.org.co": "{{email_contacto_2}}",
    "quimicoinvestigacion@hptu.org.co": "{{email_contacto_3}}",
    "Coordinacioninvestigacion3@hptu.org.co": "{{email_contacto_4}}",
    "Coordinacioninvestigacion2@hptu.org.co": "{{email_contacto_5}}",
    
    # Teléfonos
    "4459123": "{{telefono_monitoreo}}",
    "3502433321": "{{telefono_servicio_tecnico}}",
    
    # Email de sistema
    "micem@netuxtecnologia.com": "{{email_sistema}}",
}

# Patrones regex para fechas en formato DD/MM/YYYY
DATE_PATTERN = re.compile(r'\b(\d{2}/\d{2}/\d{4})\b')

# Informes a procesar
INFORMES = {
    'desempeno': 'test_documents/Informe_de_calificación_de_desempeño_ejemplo_formato_actual.docx',
    'diseno': 'test_documents/Informe_de_calificación_de_diseño_ejemplo_formato_actual.docx',
    'instalacion': 'test_documents/Informe_de_calificación_de_instalación_ejemplo_formato_actual.docx',
    'operacion': 'test_documents/Informe_de_calificación_de_operacion_ejemplo_formato_actual.docx',
}


def replace_in_paragraph(paragraph, replacements):
    """
    Reemplaza texto en un párrafo preservando formato.
    
    Args:
        paragraph: Párrafo de python-docx
        replacements: Dict con texto a reemplazar
        
    Returns:
        Número de reemplazos realizados
    """
    count = 0
    for run in paragraph.runs:
        original_text = run.text
        new_text = original_text
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
                count += 1
        if new_text != original_text:
            run.text = new_text
    return count


def replace_in_table(table, replacements):
    """
    Reemplaza texto en todas las celdas de una tabla.
    
    Args:
        table: Tabla de python-docx
        replacements: Dict con texto a reemplazar
        
    Returns:
        Número de reemplazos realizados
    """
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count += replace_in_paragraph(paragraph, replacements)
    return count


def find_placeholders_in_doc(doc):
    """
    Encuentra todos los placeholders {{...}} en el documento.
    
    Args:
        doc: Documento python-docx
        
    Returns:
        Set de placeholders encontrados
    """
    placeholders = set()
    pattern = re.compile(r'\{\{([a-zA-Z0-9_]+)\}\}')
    
    # Body
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        placeholders.update(matches)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = pattern.findall(para.text)
                    placeholders.update(matches)
    
    # Headers & Footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                matches = pattern.findall(para.text)
                placeholders.update(matches)
        if section.footer:
            for para in section.footer.paragraphs:
                matches = pattern.findall(para.text)
                placeholders.update(matches)
    
    return placeholders


def get_image_info(doc):
    """
    Obtiene información de imágenes del documento.
    
    Args:
        doc: Documento python-docx
        
    Returns:
        Dict con conteo de imágenes por ubicación
    """
    replacer = ImageReplacer(doc)
    summary = replacer.get_summary()
    return {
        'header_images': summary['total_headers'],
        'footer_images': summary['total_footers'],
        'body_images': summary['total_body'],
        'total_images': summary['total']
    }


def process_document(input_path, output_path, replacements):
    """
    Procesa un documento y crea la plantilla.
    
    Args:
        input_path: Ruta al documento original
        output_path: Ruta para guardar la plantilla
        replacements: Dict con texto a reemplazar
        
    Returns:
        Dict con información de la plantilla creada
    """
    print(f"\n📄 Procesando: {Path(input_path).name}")
    
    doc = Document(input_path)
    total_replacements = 0
    
    # Reemplazar en párrafos del cuerpo
    for para in doc.paragraphs:
        total_replacements += replace_in_paragraph(para, replacements)
    
    # Reemplazar en tablas
    for table in doc.tables:
        total_replacements += replace_in_table(table, replacements)
    
    # Reemplazar en headers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                total_replacements += replace_in_paragraph(para, replacements)
        if section.footer:
            for para in section.footer.paragraphs:
                total_replacements += replace_in_paragraph(para, replacements)
    
    # Guardar plantilla
    doc.save(output_path)
    
    # Recargar para obtener info
    doc = Document(output_path)
    placeholders = find_placeholders_in_doc(doc)
    image_info = get_image_info(doc)
    
    print(f"   ✓ Reemplazos realizados: {total_replacements}")
    print(f"   ✓ Placeholders: {len(placeholders)}")
    print(f"   ✓ Imágenes: {image_info['total_images']}")
    print(f"   ✓ Plantilla guardada: {output_path}")
    
    return {
        'replacements_made': total_replacements,
        'placeholders': sorted(list(placeholders)),
        'images': image_info
    }


def main():
    """Función principal que procesa todos los informes."""
    print("="*60)
    print("CREANDO PLANTILLAS CON PLACEHOLDERS")
    print("="*60)
    
    # Crear carpeta para plantillas
    templates_dir = Path('templates')
    templates_dir.mkdir(exist_ok=True)
    
    # Configuración de todas las plantillas
    config = {
        'version': '1.0',
        'description': 'Configuración de plantillas de informes de calificación',
        'placeholders_globales': sorted(list(set(REPLACEMENTS.values()))),
        'plantillas': {}
    }
    
    total = 0
    for nombre, input_path in INFORMES.items():
        if not Path(input_path).exists():
            print(f"\n⚠️  Archivo no encontrado: {input_path}")
            continue
            
        output_path = templates_dir / f"plantilla_{nombre}.docx"
        info = process_document(input_path, str(output_path), REPLACEMENTS)
        total += info['replacements_made']
        
        config['plantillas'][nombre] = {
            'archivo': f"plantilla_{nombre}.docx",
            'placeholders': info['placeholders'],
            'imagenes': info['images'],
            'descripcion': f"Plantilla para informe de calificación de {nombre}"
        }
    
    # Guardar configuración JSON
    config_path = templates_dir / 'plantilla_config.json'
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)
    
    print("\n" + "="*60)
    print(f"✅ COMPLETADO: {len(config['plantillas'])} plantillas creadas")
    print(f"   Total de reemplazos: {total}")
    print(f"   Ubicación: {templates_dir.absolute()}")
    print(f"   Configuración: {config_path}")
    print("="*60)
    
    # Mostrar placeholders utilizados
    print("\n📋 PLACEHOLDERS DISPONIBLES:")
    for ph in sorted(set(REPLACEMENTS.values())):
        print(f"   {ph}")
    
    # Mostrar resumen de imágenes
    print("\n🖼️  RESUMEN DE IMÁGENES POR PLANTILLA:")
    for nombre, info in config['plantillas'].items():
        imgs = info['imagenes']
        print(f"   {nombre}: {imgs['total_images']} total "
              f"(header: {imgs['header_images']}, "
              f"body: {imgs['body_images']}, "
              f"footer: {imgs['footer_images']})")


if __name__ == '__main__':
    main()
