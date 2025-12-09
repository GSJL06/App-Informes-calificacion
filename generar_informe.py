#!/usr/bin/env python
"""
Script para generar informes a partir de plantillas.
Reemplaza placeholders de texto e imágenes para crear documentos finales.
Soporta listas dinámicas y tablas con filas variables.

Uso:
    python generar_informe.py --plantilla templates/plantilla_desempeno.docx --datos datos.json --output output.docx
    python generar_informe.py --plantilla templates/plantilla_diseno.docx --datos ejemplo_datos.json --imagenes test_images/ --output informe_final.docx
"""
import sys
sys.path.insert(0, 'src')

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Any, Optional
from copy import deepcopy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from core.placeholder_engine import PlaceholderEngine
from core.image_replacer import ImageReplacer
import logging

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def load_json_data(json_path: str) -> dict:
    """
    Carga datos de reemplazo desde un archivo JSON.
    
    Args:
        json_path: Ruta al archivo JSON
        
    Returns:
        Dict con los datos de reemplazo
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def find_images_in_folder(folder_path: str) -> dict:
    """
    Busca imágenes en una carpeta y las mapea para reemplazo.
    
    Espera nombres de archivo con formato:
    - header_0_0.png (primera imagen del header de la primera sección)
    - body_0.png (primera imagen del cuerpo)
    - footer_0_0.png (primera imagen del footer)
    
    Args:
        folder_path: Ruta a la carpeta con imágenes
        
    Returns:
        Dict con mapeo de ubicación a ruta de imagen
    """
    folder = Path(folder_path)
    if not folder.exists():
        logger.warning(f"Carpeta de imágenes no encontrada: {folder_path}")
        return {}
    
    image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}
    replacements = {}
    
    for img_file in folder.iterdir():
        if img_file.suffix.lower() in image_extensions:
            # Extraer nombre sin extensión como clave
            key = img_file.stem  # ej: "header_0_0", "body_5"
            replacements[key] = str(img_file)
            logger.info(f"Imagen encontrada: {key} -> {img_file}")
    
    return replacements


class DynamicContentProcessor:
    """Processes dynamic lists and tables in DOCX documents."""
    
    PLACEHOLDER_PATTERN = r'\{\{([a-zA-Z0-9_]+)\}\}'
    LIST_PLACEHOLDER_PATTERN = r'\{\{(lista_[a-zA-Z0-9_]+)\}\}'
    TABLE_ROW_PATTERN = r'\{\{(fila_[a-zA-Z0-9_]+)\}\}'
    
    def __init__(self, document: Document):
        self.document = document
        self.pattern = re.compile(self.PLACEHOLDER_PATTERN)
        
    def expand_dynamic_lists(self, data: Dict[str, Any]) -> int:
        """
        Expand list placeholders into bullet lists.
        
        Args:
            data: Dict where list values are arrays of strings
            
        Returns:
            Number of lists expanded
        """
        count = 0
        paragraphs_to_process = []
        
        # Find paragraphs with list placeholders
        for idx, para in enumerate(self.document.paragraphs):
            for key, value in data.items():
                if isinstance(value, list) and f"{{{{{key}}}}}" in para.text:
                    paragraphs_to_process.append((idx, para, key, value))
        
        # Process in reverse to maintain indices
        for idx, para, key, items in reversed(paragraphs_to_process):
            if not items:
                para.text = para.text.replace(f"{{{{{key}}}}}", "")
                continue
            
            # Get original formatting
            original_style = para.style
            original_format = self._get_paragraph_format(para)
            
            # Replace placeholder with first item
            first_item = str(items[0])
            para.clear()
            run = para.add_run(f"• {first_item}")
            
            # Add remaining items as new paragraphs
            parent = para._element.getparent()
            para_index = list(parent).index(para._element)
            
            for item in items[1:]:
                new_para = self._create_bullet_paragraph(str(item), original_style)
                parent.insert(para_index + 1, new_para._element)
                para_index += 1
            
            count += 1
            logger.info(f"Expanded list '{key}' with {len(items)} items")
        
        return count
    
    def expand_dynamic_tables(self, data: Dict[str, Any]) -> int:
        """
        Expand table rows from array data.
        
        Args:
            data: Dict where table values are arrays of dicts
            
        Returns:
            Number of tables expanded
        """
        count = 0
        
        for table in self.document.tables:
            # Check if table has row template placeholder
            template_row_idx = None
            template_row_data = None
            placeholder_key = None
            
            for row_idx, row in enumerate(table.rows):
                row_text = ''.join(cell.text for cell in row.cells)
                
                # Look for array placeholders in row
                for key, value in data.items():
                    if isinstance(value, list) and len(value) > 0 and isinstance(value[0], dict):
                        # Check if any cell references this array's fields
                        for cell in row.cells:
                            if f"{{{{{key}}}}}" in cell.text or any(
                                f"{{{{{key}.{field}}}}}" in cell.text
                                for field in value[0].keys()
                            ):
                                template_row_idx = row_idx
                                template_row_data = value
                                placeholder_key = key
                                break
                
                if template_row_idx is not None:
                    break
            
            if template_row_idx is not None and template_row_data:
                self._expand_table_rows(table, template_row_idx, template_row_data, placeholder_key)
                count += 1
        
        return count
    
    def _expand_table_rows(self, table: Table, template_row_idx: int,
                           rows_data: List[Dict], key: str):
        """Expand a table with multiple rows from data."""
        if not rows_data:
            return
        
        template_row = table.rows[template_row_idx]
        
        # Get cell formats from template
        cell_formats = []
        for cell in template_row.cells:
            cell_formats.append({
                'text': cell.text,
                'width': cell.width if hasattr(cell, 'width') else None
            })
        
        # Clear template row and fill with first data item
        self._fill_row_with_data(template_row, cell_formats, rows_data[0], key)
        
        # Add additional rows
        for row_data in rows_data[1:]:
            new_row = self._add_table_row(table, template_row_idx + 1)
            self._fill_row_with_data(new_row, cell_formats, row_data, key)
        
        logger.info(f"Expanded table with {len(rows_data)} rows for '{key}'")
    
    def _fill_row_with_data(self, row, cell_formats: List[Dict],
                            row_data: Dict, key: str):
        """Fill a table row with data."""
        for idx, cell in enumerate(row.cells):
            if idx < len(cell_formats):
                template_text = cell_formats[idx]['text']
                new_text = template_text
                
                # Replace field placeholders
                for field, value in row_data.items():
                    new_text = new_text.replace(f"{{{{{key}.{field}}}}}", str(value))
                    new_text = new_text.replace(f"{{{{{field}}}}}", str(value))
                
                # Clear and set new text
                for para in cell.paragraphs:
                    para.clear()
                    para.add_run(new_text)
                    break
    
    def _add_table_row(self, table: Table, position: int):
        """Add a new row to table at position."""
        tbl = table._tbl
        tr = tbl.tr_lst[position - 1] if position > 0 else tbl.tr_lst[0]
        new_tr = deepcopy(tr)
        tbl.insert(position, new_tr)
        return table.rows[position]
    
    def _get_paragraph_format(self, para: Paragraph) -> Dict:
        """Extract paragraph formatting."""
        return {
            'style': para.style,
            'alignment': para.alignment
        }
    
    def _create_bullet_paragraph(self, text: str, style) -> Paragraph:
        """Create a new bullet paragraph."""
        new_para = self.document.add_paragraph(f"• {text}")
        if style:
            new_para.style = style
        return new_para


def generate_report(
    template_path: str,
    output_path: str,
    text_data: dict = None,
    image_folder: str = None,
    image_replacements: dict = None
) -> bool:
    """
    Genera un informe a partir de una plantilla.
    
    Args:
        template_path: Ruta a la plantilla .docx
        output_path: Ruta para guardar el documento generado
        text_data: Dict con datos para reemplazar placeholders de texto
                   Soporta arrays para listas dinámicas y tablas
        image_folder: Carpeta con imágenes para reemplazo automático
        image_replacements: Dict explícito con reemplazos de imagen
        
    Returns:
        True si se generó correctamente
    """
    template_path = Path(template_path)
    if not template_path.exists():
        logger.error(f"Plantilla no encontrada: {template_path}")
        return False
    
    logger.info(f"Cargando plantilla: {template_path}")
    doc = Document(template_path)
    
    # Process dynamic content first (lists and tables)
    if text_data:
        # Separate scalar values from arrays
        scalar_data = {}
        array_data = {}
        
        for key, value in text_data.items():
            if isinstance(value, list):
                array_data[key] = value
            else:
                scalar_data[key] = value
        
        # Process dynamic lists and tables
        if array_data:
            logger.info(f"Processing {len(array_data)} dynamic content items...")
            processor = DynamicContentProcessor(doc)
            
            lists_expanded = processor.expand_dynamic_lists(array_data)
            tables_expanded = processor.expand_dynamic_tables(array_data)
            
            logger.info(f"Dynamic content: {lists_expanded} lists, {tables_expanded} tables expanded")
        
        # Replace scalar placeholders
        if scalar_data:
            logger.info(f"Reemplazando {len(scalar_data)} placeholders de texto...")
            engine = PlaceholderEngine(doc)
            
            # Validar datos
            validation = engine.validate_data(scalar_data)
            if validation['missing']:
                logger.warning(f"Placeholders sin datos: {validation['missing']}")
            
            # Realizar reemplazos
            count = engine.replace_all(scalar_data, strict=False, preserve_format=True)
            logger.info(f"Reemplazos de texto realizados: {count}")
    
    # Reemplazar imágenes
    img_replacements = image_replacements or {}
    
    if image_folder:
        folder_images = find_images_in_folder(image_folder)
        img_replacements.update(folder_images)
    
    if img_replacements:
        logger.info(f"Reemplazando {len(img_replacements)} imágenes...")
        replacer = ImageReplacer(doc)
        
        # Mostrar resumen de imágenes en plantilla
        summary = replacer.get_summary()
        logger.info(f"Imágenes en plantilla: {summary['total']} "
                   f"(headers: {summary['total_headers']}, "
                   f"body: {summary['total_body']}, "
                   f"footers: {summary['total_footers']})")
        
        results = replacer.replace_images_batch(img_replacements)
        
        success = sum(1 for v in results.values() if v)
        failed = sum(1 for v in results.values() if not v)
        logger.info(f"Imágenes reemplazadas: {success} exitosas, {failed} fallidas")
        
        if failed > 0:
            for key, result in results.items():
                if not result:
                    logger.warning(f"  - Falló: {key}")
    
    # Guardar documento
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(output_path)
    logger.info(f"Documento generado: {output_path}")
    
    return True


def main():
    """Función principal con argumentos de línea de comandos."""
    parser = argparse.ArgumentParser(
        description='Genera informes a partir de plantillas DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos:
  python generar_informe.py --plantilla templates/plantilla_desempeno.docx --datos ejemplo_datos.json --output informe.docx
  python generar_informe.py -p templates/plantilla_diseno.docx -d datos.json -i test_images/ -o output/informe_final.docx
        """
    )
    
    parser.add_argument(
        '-p', '--plantilla',
        required=True,
        help='Ruta a la plantilla .docx'
    )
    
    parser.add_argument(
        '-d', '--datos',
        help='Archivo JSON con datos para reemplazar placeholders'
    )
    
    parser.add_argument(
        '-i', '--imagenes',
        help='Carpeta con imágenes para reemplazo'
    )
    
    parser.add_argument(
        '-o', '--output',
        required=True,
        help='Ruta para guardar el documento generado'
    )
    
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Mostrar información detallada'
    )
    
    parser.add_argument(
        '--info',
        action='store_true',
        help='Mostrar información de la plantilla sin generar documento'
    )
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Modo información
    if args.info:
        show_template_info(args.plantilla)
        return
    
    # Cargar datos de texto
    text_data = None
    if args.datos:
        if not Path(args.datos).exists():
            logger.error(f"Archivo de datos no encontrado: {args.datos}")
            sys.exit(1)
        text_data = load_json_data(args.datos)
        logger.info(f"Datos cargados: {len(text_data)} campos")
    
    # Generar informe
    success = generate_report(
        template_path=args.plantilla,
        output_path=args.output,
        text_data=text_data,
        image_folder=args.imagenes
    )
    
    if success:
        print(f"\n✅ Informe generado exitosamente: {args.output}")
    else:
        print(f"\n❌ Error al generar el informe")
        sys.exit(1)


def show_template_info(template_path: str):
    """
    Muestra información detallada de una plantilla.
    
    Args:
        template_path: Ruta a la plantilla
    """
    template_path = Path(template_path)
    if not template_path.exists():
        print(f"❌ Plantilla no encontrada: {template_path}")
        return
    
    doc = Document(template_path)
    
    print(f"\n📄 INFORMACIÓN DE PLANTILLA: {template_path.name}")
    print("=" * 60)
    
    # Placeholders
    engine = PlaceholderEngine(doc)
    placeholders = engine.find_all_placeholders()
    
    print(f"\n📝 PLACEHOLDERS ({len(placeholders)}):")
    for ph in sorted(placeholders):
        print(f"   {{{{ {ph} }}}}")
    
    # Imágenes
    replacer = ImageReplacer(doc)
    summary = replacer.get_summary()
    
    print(f"\n🖼️  IMÁGENES ({summary['total']}):")
    print(f"   Headers: {summary['total_headers']}")
    print(f"   Body: {summary['total_body']}")
    print(f"   Footers: {summary['total_footers']}")
    
    if summary['total'] > 0:
        print("\n   Detalle de imágenes:")
        details = summary['details']
        
        for i, img in enumerate(details['headers']):
            print(f"   - header_{img['section_idx']}_{i}: {img.get('target', 'N/A')}")
        
        for i, img in enumerate(details['body']):
            print(f"   - body_{i}: {img.get('target', 'N/A')}")
        
        for i, img in enumerate(details['footers']):
            print(f"   - footer_{img['section_idx']}_{i}: {img.get('target', 'N/A')}")
    
    # Estadísticas
    print(f"\n📊 ESTADÍSTICAS:")
    print(f"   Secciones: {len(doc.sections)}")
    print(f"   Párrafos: {len(doc.paragraphs)}")
    print(f"   Tablas: {len(doc.tables)}")
    
    print("\n" + "=" * 60)


if __name__ == '__main__':
    main()