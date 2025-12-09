"""
Script para crear un documento de prueba con placeholders
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_test_document():
    """Crea un documento de prueba con placeholders"""
    doc = Document()
    
    # Título
    title = doc.add_heading('Contrato de Servicios', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del contrato
    doc.add_paragraph()
    doc.add_paragraph('Fecha: {{fecha}}')
    doc.add_paragraph('Número de contrato: {{numero_contrato}}')
    doc.add_paragraph()
    
    # Partes
    doc.add_heading('Partes del Contrato', level=1)
    doc.add_paragraph('El presente contrato se celebra entre:')
    doc.add_paragraph()
    
    # Cliente
    doc.add_paragraph('CLIENTE:')
    doc.add_paragraph('Nombre: {{nombre_cliente}}')
    doc.add_paragraph('Empresa: {{empresa_cliente}}')
    doc.add_paragraph('Dirección: {{direccion_cliente}}')
    doc.add_paragraph('Email: {{email_cliente}}')
    doc.add_paragraph()
    
    # Proveedor
    doc.add_paragraph('PROVEEDOR:')
    doc.add_paragraph('Nombre: {{nombre_proveedor}}')
    doc.add_paragraph('Empresa: {{empresa_proveedor}}')
    doc.add_paragraph('Dirección: {{direccion_proveedor}}')
    doc.add_paragraph()
    
    # Servicios
    doc.add_heading('Descripción de Servicios', level=1)
    doc.add_paragraph('El proveedor se compromete a entregar los siguientes servicios:')
    doc.add_paragraph('- {{servicio_1}}')
    doc.add_paragraph('- {{servicio_2}}')
    doc.add_paragraph('- {{servicio_3}}')
    doc.add_paragraph()
    
    # Términos
    doc.add_heading('Términos y Condiciones', level=1)
    doc.add_paragraph('Monto total: ${{monto_total}}')
    doc.add_paragraph('Fecha de inicio: {{fecha_inicio}}')
    doc.add_paragraph('Fecha de finalización: {{fecha_fin}}')
    doc.add_paragraph('Forma de pago: {{forma_pago}}')
    doc.add_paragraph()
    
    # Firmas
    doc.add_heading('Firmas', level=1)
    doc.add_paragraph()
    doc.add_paragraph('_________________________          _________________________')
    doc.add_paragraph('{{nombre_cliente}}                    {{nombre_proveedor}}')
    doc.add_paragraph('Cliente                              Proveedor')
    
    # Agregar footer a la sección
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "{{empresa_proveedor}} - Documento Confidencial - Página "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Guardar
    output_path = Path('test_documents')
    output_path.mkdir(exist_ok=True)
    
    doc_path = output_path / 'contrato_template.docx'
    doc.save(str(doc_path))
    
    print(f"✓ Documento creado: {doc_path}")
    return doc_path


def create_data_file():
    """Crea un archivo JSON con datos de ejemplo"""
    import json
    
    data = {
        "fecha": "5 de Diciembre de 2024",
        "numero_contrato": "CONT-2024-001",
        "nombre_cliente": "María García López",
        "empresa_cliente": "Innovaciones Tech S.A.",
        "direccion_cliente": "Calle Principal 123, Ciudad",
        "email_cliente": "maria.garcia@innovatech.com",
        "nombre_proveedor": "Carlos Rodríguez",
        "empresa_proveedor": "Soluciones Digitales S.L.",
        "direccion_proveedor": "Avenida Central 456, Ciudad",
        "servicio_1": "Desarrollo de aplicación web",
        "servicio_2": "Mantenimiento mensual",
        "servicio_3": "Soporte técnico 24/7",
        "monto_total": "15,000.00",
        "fecha_inicio": "1 de Enero de 2025",
        "fecha_fin": "31 de Diciembre de 2025",
        "forma_pago": "Transferencia bancaria mensual"
    }
    
    output_path = Path('test_documents')
    output_path.mkdir(exist_ok=True)
    
    data_path = output_path / 'datos_contrato.json'
    with open(data_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    print(f"✓ Datos creados: {data_path}")
    return data_path


if __name__ == '__main__':
    print("Creando documento de prueba...")
    doc_path = create_test_document()
    data_path = create_data_file()
    
    print("\n" + "="*50)
    print("Archivos creados:")
    print(f"  - Documento: {doc_path}")
    print(f"  - Datos JSON: {data_path}")
    print("\nPara probar el reemplazo de placeholders:")
    print(f"  python -c \"from src.core.document_processor import DocumentProcessor; from src.core.placeholder_engine import PlaceholderEngine; import json; p=DocumentProcessor('{doc_path}'); p.load(); e=PlaceholderEngine(p.document); data=json.load(open('{data_path}', encoding='utf-8')); e.replace_all(data); p.save('test_documents/contrato_final.docx'); print('Documento procesado!')\"")

