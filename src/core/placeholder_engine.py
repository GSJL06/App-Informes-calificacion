"""
Placeholder Engine - Motor para procesamiento de variables {{key}}
Soporta reemplazo en body, headers, footers y tablas
"""
import re
from typing import Dict, List, Optional, Set, Callable
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import logging

logger = logging.getLogger(__name__)


class PlaceholderEngine:
    """Motor de procesamiento de placeholders con validación"""
    
    # Patrón para detectar placeholders {{variable}}
    PLACEHOLDER_PATTERN = r'\{\{([a-zA-Z0-9_]+)\}\}'
    
    def __init__(self, document: Document):
        """
        Args:
            document: Instancia de python-docx Document
        """
        self.document = document
        self.pattern = re.compile(self.PLACEHOLDER_PATTERN)
        
    def find_all_placeholders(self) -> Set[str]:
        """
        Encuentra todos los placeholders únicos en el documento
        
        Returns:
            Set de nombres de variables encontradas
        """
        placeholders = set()
        
        # Body paragraphs
        for para in self.document.paragraphs:
            matches = self.pattern.findall(para.text)
            placeholders.update(matches)
        
        # Tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = self.pattern.findall(para.text)
                        placeholders.update(matches)
        
        # Headers & Footers
        for section in self.document.sections:
            if section.header:
                for para in section.header.paragraphs:
                    matches = self.pattern.findall(para.text)
                    placeholders.update(matches)
            
            if section.footer:
                for para in section.footer.paragraphs:
                    matches = self.pattern.findall(para.text)
                    placeholders.update(matches)
        
        logger.info(f"Encontrados {len(placeholders)} placeholders únicos")
        return placeholders
    
    def validate_data(self, data: Dict[str, str]) -> Dict[str, List[str]]:
        """
        Valida que los datos proporcionados cubran todos los placeholders
        
        Args:
            data: Dict con valores para reemplazo
            
        Returns:
            Dict con 'missing' (placeholders sin datos) y 'unused' (datos sin uso)
        """
        found_placeholders = self.find_all_placeholders()
        provided_keys = set(data.keys())
        
        validation = {
            'missing': list(found_placeholders - provided_keys),
            'unused': list(provided_keys - found_placeholders)
        }
        
        if validation['missing']:
            logger.warning(f"Placeholders sin datos: {validation['missing']}")
        if validation['unused']:
            logger.info(f"Datos sin placeholder: {validation['unused']}")
        
        return validation
    
    def replace_all(
        self,
        data: Dict[str, str],
        strict: bool = False,
        preserve_format: bool = True
    ) -> int:
        """
        Reemplaza todos los placeholders en el documento
        
        Args:
            data: Dict con valores de reemplazo
            strict: Si True, falla si hay placeholders sin datos
            preserve_format: Mantener formato de texto
            
        Returns:
            Número total de reemplazos realizados
            
        Raises:
            ValueError: Si strict=True y hay placeholders sin datos
        """
        validation = self.validate_data(data)
        
        if strict and validation['missing']:
            raise ValueError(
                f"Placeholders sin datos: {validation['missing']}"
            )
        
        total_replacements = 0
        
        # Body
        total_replacements += self._replace_in_paragraphs(
            self.document.paragraphs, data, preserve_format
        )
        
        # Tables
        for table in self.document.tables:
            total_replacements += self._replace_in_table(
                table, data, preserve_format
            )
        
        # Headers & Footers
        for section in self.document.sections:
            if section.header:
                total_replacements += self._replace_in_paragraphs(
                    section.header.paragraphs, data, preserve_format
                )
            
            if section.footer:
                total_replacements += self._replace_in_paragraphs(
                    section.footer.paragraphs, data, preserve_format
                )
        
        logger.info(f"Total de reemplazos: {total_replacements}")
        return total_replacements
    
    def _replace_in_paragraphs(
        self,
        paragraphs: List[Paragraph],
        data: Dict[str, str],
        preserve_format: bool
    ) -> int:
        """Reemplaza placeholders en lista de párrafos"""
        count = 0
        
        for para in paragraphs:
            # Buscar placeholders en el texto completo del párrafo
            if not self.pattern.search(para.text):
                continue
            
            if preserve_format:
                # Reemplazo preservando formato de runs
                count += self._replace_in_runs(para, data)
            else:
                # Reemplazo simple del texto completo
                new_text = para.text
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in new_text:
                        new_text = new_text.replace(placeholder, str(value))
                        count += 1
                
                # Actualizar texto
                para.text = new_text
        
        return count
    
    def _replace_in_runs(self, para: Paragraph, data: Dict[str, str]) -> int:
        """Reemplaza placeholders preservando formato de runs individuales"""
        count = 0
        
        # Reconstruir el párrafo run por run
        for run in para.runs:
            if not self.pattern.search(run.text):
                continue
            
            new_text = run.text
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))
                    count += 1
            
            run.text = new_text
        
        return count
    
    def _replace_in_table(
        self,
        table: Table,
        data: Dict[str, str],
        preserve_format: bool
    ) -> int:
        """Reemplaza placeholders en una tabla"""
        count = 0
        
        for row in table.rows:
            for cell in row.cells:
                count += self._replace_in_paragraphs(
                    cell.paragraphs, data, preserve_format
                )
        
        return count
    
    def replace_with_function(
        self,
        transformer: Callable[[str], str]
    ) -> int:
        """
        Reemplaza placeholders usando una función de transformación
        
        Args:
            transformer: Función que recibe nombre de variable y retorna valor
            
        Returns:
            Número de reemplazos
            
        Example:
            def uppercase_transformer(var_name):
                return var_name.upper()
            
            engine.replace_with_function(uppercase_transformer)
        """
        placeholders = self.find_all_placeholders()
        data = {key: transformer(key) for key in placeholders}
        return self.replace_all(data, strict=False)
    
    def get_placeholder_report(self) -> Dict:
        """
        Genera reporte detallado de placeholders en el documento
        
        Returns:
            Dict con estadísticas y ubicaciones
        """
        report = {
            'total_unique': 0,
            'placeholders': {},
            'locations': {
                'body': 0,
                'tables': 0,
                'headers': 0,
                'footers': 0
            }
        }
        
        all_placeholders = self.find_all_placeholders()
        report['total_unique'] = len(all_placeholders)
        
        # Inicializar contadores
        for ph in all_placeholders:
            report['placeholders'][ph] = 0
        
        # Contar en body
        for para in self.document.paragraphs:
            for match in self.pattern.finditer(para.text):
                var_name = match.group(1)
                report['placeholders'][var_name] += 1
                report['locations']['body'] += 1
        
        # Contar en tablas
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for match in self.pattern.finditer(para.text):
                            var_name = match.group(1)
                            report['placeholders'][var_name] += 1
                            report['locations']['tables'] += 1
        
        # Contar en headers/footers
        for section in self.document.sections:
            if section.header:
                for para in section.header.paragraphs:
                    for match in self.pattern.finditer(para.text):
                        var_name = match.group(1)
                        report['placeholders'][var_name] += 1
                        report['locations']['headers'] += 1
            
            if section.footer:
                for para in section.footer.paragraphs:
                    for match in self.pattern.finditer(para.text):
                        var_name = match.group(1)
                        report['placeholders'][var_name] += 1
                        report['locations']['footers'] += 1
        
        return report
    
    def preview_replacements(
        self,
        data: Dict[str, str],
        max_examples: int = 5
    ) -> List[Dict[str, str]]:
        """
        Vista previa de reemplazos sin modificar documento
        
        Args:
            data: Dict con valores de reemplazo
            max_examples: Máximo de ejemplos a retornar
            
        Returns:
            Lista de dicts con 'original' y 'replaced'
        """
        examples = []
        count = 0
        
        for para in self.document.paragraphs:
            if count >= max_examples:
                break
            
            if self.pattern.search(para.text):
                original = para.text
                replaced = original
                
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    replaced = replaced.replace(placeholder, str(value))
                
                if original != replaced:
                    examples.append({
                        'original': original,
                        'replaced': replaced
                    })
                    count += 1
        
        return examples