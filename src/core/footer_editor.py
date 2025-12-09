"""
Footer Editor - Editor de pies de página con preservación de formato
Soporta múltiples secciones, formatos y numeración de páginas
"""
from typing import Dict, List, Optional, Tuple, Union
from docx import Document
from docx.section import Section
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging

logger = logging.getLogger(__name__)


class FooterEditor:
    """Editor de pies de página con preservación completa de formato"""

    def __init__(self, document: Document):
        """
        Args:
            document: Instancia de python-docx Document
        """
        self.document = document

    def get_sections_count(self) -> int:
        """Retorna el número de secciones en el documento"""
        return len(self.document.sections)

    def get_footer(self, section_idx: int = 0) -> Optional[object]:
        """
        Obtiene el footer de una sección específica

        Args:
            section_idx: Índice de la sección (default: 0)

        Returns:
            Objeto footer o None si no existe
        """
        sections = self.document.sections
        if section_idx >= len(sections):
            raise IndexError(f"Sección {section_idx} no existe. Total: {len(sections)}")

        return sections[section_idx].footer

    def get_footer_text(self, section_idx: int = 0) -> str:
        """
        Obtiene el texto del footer de una sección

        Args:
            section_idx: Índice de la sección

        Returns:
            Texto completo del footer
        """
        footer = self.get_footer(section_idx)
        if footer is None:
            return ""

        text_parts = []
        for paragraph in footer.paragraphs:
            text_parts.append(paragraph.text)

        return '\n'.join(text_parts)

    def get_footer_with_format(self, section_idx: int = 0) -> List[Dict]:
        """
        Obtiene el footer con información de formato

        Args:
            section_idx: Índice de la sección

        Returns:
            Lista de diccionarios con texto y formato de cada párrafo
        """
        footer = self.get_footer(section_idx)
        if footer is None:
            return []

        result = []
        for paragraph in footer.paragraphs:
            para_info = {
                'text': paragraph.text,
                'alignment': str(paragraph.alignment) if paragraph.alignment else 'LEFT',
                'runs': []
            }

            for run in paragraph.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': str(run.font.size) if run.font.size else None,
                    'color': self._get_run_color(run)
                }
                para_info['runs'].append(run_info)

            result.append(para_info)

        return result

    def _get_run_color(self, run: Run) -> Optional[str]:
        """Extrae el color de un run como string hex"""
        try:
            if run.font.color and run.font.color.rgb:
                return str(run.font.color.rgb)
        except Exception:
            pass
        return None

    def update_footer_text(
        self,
        text: str,
        section_idx: int = 0,
        preserve_format: bool = True,
        alignment: Optional[WD_ALIGN_PARAGRAPH] = None
    ) -> None:
        """
        Actualiza el texto del footer preservando formato

        Args:
            text: Nuevo texto para el footer
            section_idx: Índice de la sección
            preserve_format: Mantener formato del primer run existente
            alignment: Alineación del párrafo (opcional)
        """
        footer = self.get_footer(section_idx)

        # Guardar formato del primer run si existe
        saved_format = None
        if preserve_format and footer.paragraphs:
            for para in footer.paragraphs:
                if para.runs:
                    saved_format = self._extract_run_format(para.runs[0])
                    break

        # Limpiar footer existente
        for paragraph in footer.paragraphs:
            paragraph.clear()

        # Si no hay párrafos, usar el primero o crear uno
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()

        # Agregar nuevo texto
        run = paragraph.add_run(text)

        # Aplicar formato guardado
        if saved_format:
            self._apply_run_format(run, saved_format)

        # Aplicar alineación
        if alignment:
            paragraph.alignment = alignment

        logger.info(f"Footer actualizado en sección {section_idx}")

    def _extract_run_format(self, run: Run) -> Dict:
        """Extrae el formato de un run"""
        return {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'color_rgb': run.font.color.rgb if run.font.color else None
        }

    def _apply_run_format(self, run: Run, format_dict: Dict) -> None:
        """Aplica formato a un run"""
        if format_dict.get('bold') is not None:
            run.bold = format_dict['bold']
        if format_dict.get('italic') is not None:
            run.italic = format_dict['italic']
        if format_dict.get('underline') is not None:
            run.underline = format_dict['underline']
        if format_dict.get('font_name'):
            run.font.name = format_dict['font_name']
        if format_dict.get('font_size'):
            run.font.size = format_dict['font_size']
        if format_dict.get('color_rgb'):
            run.font.color.rgb = format_dict['color_rgb']

    def update_footer_formatted(
        self,
        text_parts: List[Dict],
        section_idx: int = 0,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER
    ) -> None:
        """
        Actualiza footer con múltiples formatos

        Args:
            text_parts: Lista de dicts con 'text', 'bold', 'italic', 'font_name',
                       'font_size', 'color' (tuple RGB)
            section_idx: Índice de la sección
            alignment: Alineación del párrafo
        """
        footer = self.get_footer(section_idx)

        # Limpiar footer
        for paragraph in footer.paragraphs:
            paragraph.clear()

        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
        else:
            paragraph = footer.add_paragraph()

        paragraph.alignment = alignment

        for part in text_parts:
            run = paragraph.add_run(part.get('text', ''))

            if part.get('bold'):
                run.bold = True
            if part.get('italic'):
                run.italic = True
            if part.get('font_name'):
                run.font.name = part['font_name']
            if part.get('font_size'):
                run.font.size = Pt(part['font_size'])
            if part.get('color'):
                r, g, b = part['color']
                run.font.color.rgb = RGBColor(r, g, b)

        logger.info(f"Footer con formato múltiple aplicado en sección {section_idx}")

    def apply_to_all_sections(self, text: str, preserve_format: bool = True) -> int:
        """
        Aplica el mismo footer a todas las secciones

        Args:
            text: Texto para el footer
            preserve_format: Preservar formato existente

        Returns:
            Número de secciones actualizadas
        """
        count = 0
        for idx in range(len(self.document.sections)):
            self.update_footer_text(text, section_idx=idx, preserve_format=preserve_format)
            count += 1

        logger.info(f"Footer aplicado a {count} secciones")
        return count

    def add_page_number(
        self,
        section_idx: int = 0,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
        format_string: str = "Página {PAGE} de {NUMPAGES}"
    ) -> None:
        """
        Agrega numeración de páginas al footer

        Args:
            section_idx: Índice de la sección
            alignment: Alineación
            format_string: Formato del texto (usa {PAGE} y {NUMPAGES})
        """
        footer = self.get_footer(section_idx)

        # Crear nuevo párrafo para numeración
        paragraph = footer.add_paragraph()
        paragraph.alignment = alignment

        # Dividir el format_string y agregar campos
        parts = format_string.replace('{PAGE}', '\x00PAGE\x00').replace('{NUMPAGES}', '\x00NUMPAGES\x00').split('\x00')

        for part in parts:
            if part == 'PAGE':
                self._add_page_field(paragraph)
            elif part == 'NUMPAGES':
                self._add_numpages_field(paragraph)
            elif part:
                paragraph.add_run(part)

        logger.info(f"Numeración de páginas agregada en sección {section_idx}")

    def _add_page_field(self, paragraph) -> None:
        """Agrega campo PAGE al párrafo"""
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.text = "PAGE"

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

    def _add_numpages_field(self, paragraph) -> None:
        """Agrega campo NUMPAGES al párrafo"""
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.text = "NUMPAGES"

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

    def clear_footer(self, section_idx: int = 0) -> None:
        """
        Limpia completamente el footer de una sección

        Args:
            section_idx: Índice de la sección
        """
        footer = self.get_footer(section_idx)
        for paragraph in footer.paragraphs:
            paragraph.clear()

        logger.info(f"Footer limpiado en sección {section_idx}")
