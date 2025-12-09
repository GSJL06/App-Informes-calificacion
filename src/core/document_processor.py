"""
Core Document Processor - Motor principal para edición OOXML
Optimizado para archivos hasta 20MB con preservación de formato
"""
import os
import zipfile
import shutil
from pathlib import Path
from typing import Dict, Optional, List, Union
from datetime import datetime
from lxml import etree
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Procesador principal de documentos DOCX con enfoque en rendimiento"""
    
    # Namespaces OOXML
    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }
    
    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
    
    def __init__(self, file_path: Union[str, Path]):
        """
        Inicializa el procesador con validación de archivo
        
        Args:
            file_path: Ruta al archivo .docx
            
        Raises:
            FileNotFoundError: Si el archivo no existe
            ValueError: Si el archivo excede límite de tamaño
        """
        self.file_path = Path(file_path)
        self._validate_file()
        self.document = None
        self._backup_path = None
        
    def _validate_file(self) -> None:
        """Valida existencia y tamaño del archivo"""
        if not self.file_path.exists():
            raise FileNotFoundError(f"Archivo no encontrado: {self.file_path}")
        
        file_size = self.file_path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            raise ValueError(
                f"Archivo excede límite de {self.MAX_FILE_SIZE/1024/1024}MB: "
                f"{file_size/1024/1024:.2f}MB"
            )
        
        if not zipfile.is_zipfile(self.file_path):
            raise ValueError("Archivo no es un documento .docx válido")
    
    def load(self) -> 'DocumentProcessor':
        """Carga el documento en memoria"""
        try:
            logger.info(f"Cargando documento: {self.file_path}")
            self.document = Document(self.file_path)
            logger.debug(f"Documento cargado: {len(self.document.paragraphs)} párrafos")
            return self
        except Exception as e:
            logger.error(f"Error al cargar documento: {e}")
            raise
    
    def create_backup(self, backup_dir: Optional[Path] = None) -> Path:
        """
        Crea backup timestamped del documento original
        
        Args:
            backup_dir: Directorio para backups (default: mismo directorio)
            
        Returns:
            Path del backup creado
        """
        if backup_dir is None:
            backup_dir = self.file_path.parent
        else:
            backup_dir = Path(backup_dir)
            backup_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{self.file_path.stem}.backup.{timestamp}{self.file_path.suffix}"
        backup_path = backup_dir / backup_name
        
        shutil.copy2(self.file_path, backup_path)
        self._backup_path = backup_path
        logger.info(f"Backup creado: {backup_path}")
        
        return backup_path
    
    def save(self, output_path: Optional[Union[str, Path]] = None) -> Path:
        """
        Guarda el documento modificado
        
        Args:
            output_path: Ruta de salida (default: sobrescribe original)
            
        Returns:
            Path del archivo guardado
        """
        if self.document is None:
            raise RuntimeError("Documento no cargado. Ejecutar load() primero")
        
        if output_path is None:
            output_path = self.file_path
        else:
            output_path = Path(output_path)
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            self.document.save(output_path)
            logger.info(f"Documento guardado: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error al guardar documento: {e}")
            raise
    
    def get_sections(self) -> List:
        """Obtiene todas las secciones del documento"""
        if self.document is None:
            raise RuntimeError("Documento no cargado")
        return self.document.sections
    
    def get_core_properties(self) -> Dict[str, str]:
        """Extrae propiedades del documento (metadatos)"""
        if self.document is None:
            raise RuntimeError("Documento no cargado")
        
        props = self.document.core_properties
        return {
            'title': props.title or '',
            'author': props.author or '',
            'subject': props.subject or '',
            'keywords': props.keywords or '',
            'created': str(props.created) if props.created else '',
            'modified': str(props.modified) if props.modified else '',
        }
    
    def update_core_properties(self, properties: Dict[str, str]) -> None:
        """
        Actualiza propiedades del documento
        
        Args:
            properties: Dict con propiedades a actualizar
        """
        if self.document is None:
            raise RuntimeError("Documento no cargado")
        
        props = self.document.core_properties
        
        for key, value in properties.items():
            if hasattr(props, key):
                setattr(props, key, value)
                logger.debug(f"Propiedad actualizada: {key} = {value}")
    
    def get_statistics(self) -> Dict[str, int]:
        """Obtiene estadísticas del documento"""
        if self.document is None:
            raise RuntimeError("Documento no cargado")
        
        return {
            'paragraphs': len(self.document.paragraphs),
            'sections': len(self.document.sections),
            'tables': len(self.document.tables),
            'file_size_bytes': self.file_path.stat().st_size,
        }
    
    def extract_text(self, include_headers_footers: bool = False) -> str:
        """
        Extrae todo el texto del documento
        
        Args:
            include_headers_footers: Incluir encabezados/pies de página
            
        Returns:
            Texto completo del documento
        """
        if self.document is None:
            raise RuntimeError("Documento no cargado")
        
        text_parts = []
        
        # Texto principal
        for para in self.document.paragraphs:
            text_parts.append(para.text)
        
        # Headers y footers
        if include_headers_footers:
            for section in self.document.sections:
                # Headers
                if section.header:
                    for para in section.header.paragraphs:
                        text_parts.append(para.text)
                # Footers
                if section.footer:
                    for para in section.footer.paragraphs:
                        text_parts.append(para.text)
        
        return '\n'.join(text_parts)
    
    def validate_integrity(self) -> Dict[str, bool]:
        """
        Valida integridad del documento OOXML
        
        Returns:
            Dict con resultados de validación
        """
        results = {
            'is_valid_zip': False,
            'has_document_xml': False,
            'has_rels': False,
            'xml_well_formed': False
        }
        
        try:
            # Validar ZIP
            with zipfile.ZipFile(self.file_path, 'r') as zf:
                results['is_valid_zip'] = True
                
                # Validar estructura
                namelist = zf.namelist()
                results['has_document_xml'] = 'word/document.xml' in namelist
                results['has_rels'] = '_rels/.rels' in namelist
                
                # Validar XML
                if results['has_document_xml']:
                    xml_content = zf.read('word/document.xml')
                    etree.fromstring(xml_content)
                    results['xml_well_formed'] = True
        
        except Exception as e:
            logger.error(f"Error en validación de integridad: {e}")
        
        return results


class PerformanceMonitor:
    """Monitor de rendimiento para operaciones del documento"""
    
    def __init__(self):
        self.metrics = {}
        self.start_time = None
    
    def start(self, operation: str):
        """Inicia medición de operación"""
        self.start_time = datetime.now()
        logger.debug(f"Iniciando operación: {operation}")
    
    def end(self, operation: str):
        """Finaliza medición y registra métricas"""
        if self.start_time is None:
            return
        
        elapsed = (datetime.now() - self.start_time).total_seconds()
        self.metrics[operation] = elapsed
        logger.info(f"Operación '{operation}' completada en {elapsed:.3f}s")
        self.start_time = None
    
    def get_metrics(self) -> Dict[str, float]:
        """Obtiene todas las métricas registradas"""
        return self.metrics.copy()
    
    def reset(self):
        """Reinicia métricas"""
        self.metrics.clear()
        self.start_time = None